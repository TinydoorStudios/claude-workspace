#!/usr/bin/env python3
"""
Advance docx service — n8n POSTs a form submission, this fills Brian's internal
production advance sheet (docx template) and drops it into Dropbox.

Model: one tagged template PER venue/series, each a copy of Brian's real blank master
with the band-provided cells tagged. Everything else — schedule, PA, consoles, crew,
emcee, "Salsa on the Square" — stays baked exactly as his master has it. The service
only injects what the band submits. Add a new venue = tag its blank master and register
it in TEMPLATES; no code defaults to keep in sync.

Env (see advance.env.example):
  ADVANCE_PORT   default 8097
  ADVANCE_SECRET shared secret; must match n8n + Apps Script
  DROPBOX_TOKEN  Dropbox app access token (files.content.write + sharing.write)
  DROPBOX_BASE   Dropbox target folder, e.g. /3CDC Advancing
  UPLOAD_EMAIL   the dedicated stage-plot inbox (shown in the Stage Plot cell)
"""

import io
import os
import re
from flask import Flask, request, jsonify

app = Flask(__name__)

PORT = int(os.environ.get("ADVANCE_PORT", "8097"))
SECRET = os.environ.get("ADVANCE_SECRET", "")
DROPBOX_TOKEN = os.environ.get("DROPBOX_TOKEN", "")
DROPBOX_BASE = os.environ.get("DROPBOX_BASE", "/3CDC Advancing")
UPLOAD_EMAIL = os.environ.get("UPLOAD_EMAIL", "the advance inbox")
TPL_DIR = os.path.join(os.path.dirname(__file__), "templates")

# venue → tagged blank master. Add rows as Brian supplies each venue's blank doc.
TEMPLATES = {
    "Fountain Square": "advance_fsq_salsa.docx",
}


def g(fields, label, default=""):
    return fields.get(label, default) or default


def build_context(payload):
    """Only the band-provided cells; every default lives in the template itself."""
    f = payload.get("fields", {}) or {}

    notes = []
    if g(f, "Are you bringing your own engineer?") == "Yes":
        notes.append("Own engineer: " + (g(f, "If yes — engineer name & what they’re mixing (FOH/MON)") or "Yes"))
    if g(f, "Any special guest performers?") == "Yes":
        notes.append("Special guests: " + (g(f, "If yes — who, and doing what?") or "Yes"))

    stageplot = "Emailed to %s" % UPLOAD_EMAIL
    link = g(f, "Stage plot / input list — link (optional)")
    if link:
        stageplot += "  ·  " + link

    lv = ""
    if g(f, "Do you need large-vehicle parking?") == "Yes":
        lv = "— Large vehicle: " + (g(f, "If yes — vehicle type / size") or "yes")

    return {
        "band": payload.get("act") or g(f, "Act / band name"),
        "show_date": payload.get("date") or g(f, "Show date"),
        "performer_name": g(f, "Advancing contact — your name"),
        "performer_phone": g(f, "Best phone to reach the band day-of"),
        "performer_email": payload.get("email") or f.get("__email", ""),
        "monitors": g(f, "Monitor needs"),
        "input_notes": "; ".join(notes),
        "merch": g(f, "Are you selling merch?"),
        "dressing_tent": g(f, "Would you like a 10x10 private tent with sidewalls?"),
        "performers": g(f, "Total number of performers"),
        "escort_rep": g(f, "Band representative with stage-escort ability — name"),
        "wristbands": g(f, "Total wristbands needed"),
        "large_vehicle_note": lv,
        "stageplot": stageplot,
    }


def template_for(venue):
    name = TEMPLATES.get(venue)
    if not name:
        return None
    path = os.path.join(TPL_DIR, name)
    return path if os.path.exists(path) else None


def render_docx(ctx, venue):
    tpl_path = template_for(venue)
    if tpl_path:
        from docxtpl import DocxTemplate
        tpl = DocxTemplate(tpl_path)
        tpl.render(ctx)
        buf = io.BytesIO()
        tpl.save(buf)
        return buf.getvalue()

    # No template for this venue yet — emit a plain data doc so nothing is lost.
    from docx import Document
    doc = Document()
    doc.add_heading("3CDC Advance — %s" % (ctx.get("band") or "Unknown"), level=1)
    doc.add_paragraph("Venue: %s   Date: %s   (no template registered for this venue yet)"
                      % (venue or "—", ctx.get("show_date") or "—"))
    for k, v in ctx.items():
        doc.add_paragraph("%s: %s" % (k, v if v else "—"))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def safe(s):
    return re.sub(r"[^A-Za-z0-9 _.-]", "", str(s or "")).strip() or "Unknown"


def upload_dropbox(data, ctx, venue):
    fname = "Advance - %s - %s.docx" % (safe(ctx.get("band")), safe(ctx.get("show_date")))
    path = "%s/%s/%s" % (DROPBOX_BASE.rstrip("/"), safe(venue), fname)
    if not DROPBOX_TOKEN:
        return None, path
    import dropbox
    dbx = dropbox.Dropbox(DROPBOX_TOKEN)
    dbx.files_upload(data, path, mode=dropbox.files.WriteMode.overwrite)
    try:
        link = dbx.sharing_create_shared_link_with_settings(path).url
    except dropbox.exceptions.ApiError:
        links = dbx.sharing_list_shared_links(path=path).links
        link = links[0].url if links else None
    return link, path


@app.route("/health")
def health():
    return jsonify(ok=True, venues=list(TEMPLATES.keys()))


@app.route("/render", methods=["POST"])
def render():
    payload = request.get_json(force=True, silent=True) or {}
    if SECRET and payload.get("secret") != SECRET:
        return jsonify(error="bad secret"), 403
    venue = payload.get("venue") or (payload.get("fields", {}) or {}).get("Venue", "")
    ctx = build_context(payload)
    data = render_docx(ctx, venue)
    link, path = upload_dropbox(data, ctx, venue)
    return jsonify(ok=True, dropbox_path=path, dropbox_link=link,
                   act=ctx.get("band"), venue=venue, show_date=ctx.get("show_date"),
                   template=bool(template_for(venue)))


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=PORT)
