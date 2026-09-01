#!/usr/bin/env python3
"""Doc-fill engine — turn an advance submission into a filled Word document.

This replaces the manual step of copying an artist's email answers into the
standard show DOC. It renders a .docx template (docxtpl / Jinja placeholders)
using the artist's newest submission from the database.

  # one-time: generate a stand-in template you can replace with the real 3CDC DOC
  python3 docfill.py --make-sample

  # fill for a band (uses their newest submission)
  python3 docfill.py --artist "Buffalo Wabs and the Price Hill Hustle"

  # or by submission id
  python3 docfill.py --submission 12

  # list every placeholder the template can use
  python3 docfill.py --fields

Drop your real template at tools/doc_templates/advance_sheet.docx using the same
{{ placeholder }} names (see --fields). Output lands in tools/filled/.
"""
import argparse
import datetime as dt
import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent
for _cand in (HERE.parent, HERE.parent / "app"):  # deployed flat, or repo layout
    if (_cand / "advance_db.py").exists():
        sys.path.insert(0, str(_cand))
        break
import advance_db as db

TEMPLATES = HERE / "doc_templates"
FILLED = HERE / "filled"
TEMPLATES.mkdir(exist_ok=True)
FILLED.mkdir(exist_ok=True)
DEFAULT_TEMPLATE = TEMPLATES / "advance_sheet.docx"
SAMPLE_TEMPLATE = TEMPLATES / "advance_sheet_SAMPLE.docx"

FIELDS = [
    ("band_name", "Exact band name"),
    ("venue", "Venue"),
    ("show_date", "Show date (YYYY-MM-DD)"),
    ("show_series", "Show series key"),
    ("contact_name", "Day-of contact name (if collected)"),
    ("contact_phone", "Day-of contact phone"),
    ("contact_email", "Contact email (from advance list)"),
    ("performers", "Total performers + crew"),
    ("monitors", "Monitor count"),
    ("own_iems", "Bringing own IEMs (Yes/No)"),
    ("split_snake", "Providing split snake (Yes/No, if IEMs)"),
    ("stage_type", "Flat stage / Drum riser"),
    ("own_engineer", "Bringing own engineer"),
    ("merch", "Selling merch (Yes/No)"),
    ("band_tent", "Private band tent"),
    ("large_vehicle", "Needs large-vehicle parking (Yes/No)"),
    ("backline", "Backline / sharing notes"),
    ("scenic", "Backdrop / scenic elements"),
    ("lighting", "Lighting requests"),
    ("stage_plot_desc", "Stage plot / input list description or link"),
    ("stage_plot_file", "Uploaded stage plot filename (if any)"),
    ("additional", "Additional questions or concerns"),
    ("submitted_at", "When the artist submitted"),
    ("generated_on", "When this document was generated"),
]


def yn(v):
    return "Yes" if v is True else ("No" if v is False else "")


def build_context(sub, artist):
    d = sub.get("data") or {}
    def g(key):
        # promoted column first, fall back to raw form data
        return sub.get(key) if sub.get(key) not in (None, "") else d.get(key, "")
    ctx = {
        "band_name": artist["name"],
        "venue": g("venue"),
        "show_date": sub["show_date"].isoformat() if sub.get("show_date") else d.get("show_date", ""),
        "show_series": d.get("show_series", ""),
        "contact_name": g("contact_name"),
        "contact_phone": g("contact_phone"),
        "contact_email": g("contact_email"),
        "performers": g("performers"),
        "monitors": g("monitors"),
        "own_iems": yn(sub.get("own_iems")),
        "split_snake": g("split_snake"),
        "stage_type": g("stage_type"),
        "own_engineer": g("own_engineer"),
        "merch": yn(sub.get("merch")),
        "band_tent": g("band_tent"),
        "large_vehicle": yn(sub.get("large_vehicle")),
        "backline": d.get("backline", ""),
        "scenic": d.get("scenic", ""),
        "lighting": d.get("lighting", ""),
        "stage_plot_desc": d.get("stage_plot_desc", ""),
        "stage_plot_file": d.get("stage_plot_file", ""),
        "additional": d.get("additional", ""),
        "submitted_at": sub["submitted_at"].strftime("%Y-%m-%d %H:%M") if sub.get("submitted_at") else "",
        "generated_on": dt.datetime.now().strftime("%Y-%m-%d %H:%M"),
    }
    return ctx


def make_sample():
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    title = doc.add_heading("3CDC Show Advance Sheet", level=0)
    doc.add_paragraph("Generated from the band advance submission — {{ generated_on }}")

    doc.add_heading("Show", level=1)
    for label, key in [("Band", "band_name"), ("Venue", "venue"),
                       ("Date", "show_date"), ("Series", "show_series"),
                       ("Submitted", "submitted_at")]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run("{{ %s }}" % key)

    doc.add_heading("Contact", level=1)
    for label, key in [("Name", "contact_name"), ("Phone", "contact_phone"),
                       ("Email", "contact_email"),
                       ("Performers + crew", "performers")]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run("{{ %s }}" % key)

    doc.add_heading("Stage & Technical", level=1)
    for label, key in [("Stage", "stage_type"), ("Monitors", "monitors"),
                       ("Own IEMs", "own_iems"), ("Split snake", "split_snake"),
                       ("Own engineer", "own_engineer"),
                       ("Stage plot / inputs", "stage_plot_desc"),
                       ("Stage plot file", "stage_plot_file"),
                       ("Backline / sharing", "backline"),
                       ("Scenic", "scenic"), ("Lighting", "lighting")]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run("{{ %s }}" % key)

    doc.add_heading("Hospitality & Site", level=1)
    for label, key in [("Merch", "merch"), ("Band tent", "band_tent"),
                       ("Large vehicle parking", "large_vehicle")]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run("{{ %s }}" % key)

    doc.add_heading("Notes", level=1)
    p = doc.add_paragraph()
    p.add_run("Additional: ").bold = True
    p.add_run("{{ additional }}")

    SAMPLE_TEMPLATE.parent.mkdir(exist_ok=True)
    doc.save(SAMPLE_TEMPLATE)
    if not DEFAULT_TEMPLATE.exists():
        doc.save(DEFAULT_TEMPLATE)
    print(f"Sample template written to {SAMPLE_TEMPLATE}")
    print(f"Active template (replace with your real DOC): {DEFAULT_TEMPLATE}")


def fill(sub_id=None, artist_name=None):
    template = DEFAULT_TEMPLATE if DEFAULT_TEMPLATE.exists() else SAMPLE_TEMPLATE
    if not template.exists():
        print("No template found. Run: python3 docfill.py --make-sample", file=sys.stderr)
        sys.exit(1)
    from docxtpl import DocxTemplate
    with db.get_conn() as conn, conn.cursor() as cur:
        if sub_id:
            sub = db.get_submission(cur, sub_id)
            if not sub:
                print(f"No submission id {sub_id}", file=sys.stderr); sys.exit(1)
            artist = db.get_artist(cur, sub["artist_id"])
        else:
            artist = db.find_artist_by_name(cur, artist_name)
            if not artist:
                print(f"No artist matching '{artist_name}'", file=sys.stderr); sys.exit(1)
            sub = db.newest_submission(cur, artist["id"])
            if not sub:
                print(f"'{artist['name']}' has no submissions yet", file=sys.stderr); sys.exit(1)
    ctx = build_context(sub, artist)
    tpl = DocxTemplate(str(template))
    tpl.render(ctx)
    import re
    safe = re.sub(r"[^A-Za-z0-9]+", "_", artist["name"]).strip("_")
    out = FILLED / f"{safe}__{ctx['show_date'] or 'nodate'}__advance.docx"
    tpl.save(out)
    print(f"Filled document written to {out}  (template: {template.name})")
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--make-sample", action="store_true")
    ap.add_argument("--fields", action="store_true")
    ap.add_argument("--artist")
    ap.add_argument("--submission", type=int)
    args = ap.parse_args()

    if args.fields:
        print("Placeholders available to the template (use {{ name }} in the .docx):\n")
        for k, desc in FIELDS:
            print(f"  {{{{ {k} }}}}".ljust(28) + desc)
        return
    if args.make_sample:
        make_sample(); return
    if args.artist or args.submission:
        fill(sub_id=args.submission, artist_name=args.artist); return
    ap.print_help()


if __name__ == "__main__":
    main()
