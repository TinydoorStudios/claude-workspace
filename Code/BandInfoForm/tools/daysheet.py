#!/usr/bin/env python3
"""Fill an FSQ advance day-sheet — single-band, 2-band, or 3-band — from an
event. These are the only day-sheets in use (2026-09-03) — the old 513
Airwaves-specific template is retired; nothing here references it.

Template is picked automatically by how many acts the event has:
  1 act  -> doc_templates/FSQ Single Band Advance.docx
  2 acts -> doc_templates/FSQ 2 Band Advance.docx
  3 acts -> doc_templates/FSQ 3 Band Advance.docx
Acts map to columns left-to-right in event_acts' own slot_order (opener,
direct_support, headliner) — whichever slots the event actually has, in that
order, so a 2-band show booked as opener+headliner or as
direct_support+headliner both land correctly without a slot-name lookup.

Source of truth is the advance SPREADSHEET (event + any band overrides); the
advance FORM fills whatever the sheet left blank — the spreadsheet value
wins, same merge policy as the email drafts. Writes:
  - EVENT INFORMATION: Date / Event always; Band (single-band only — a
    multi-band show's names go on the act-header row instead); TONIGHT —
    MC/DJ (multi-band only)
  - Event Type / Paying?, Lead name + cell
  - per-act: Set Length, act name (multi-band header row), Stage Plot,
    Monitors, IEMs, Input Notes (also carries any lighting request / split
    snake text — there's no dedicated cell for those), Stage Type, Scenic
    Notes, Merch, Parking, Drink Tix, Dressing Room Tent, Backline, Band
    Contact — Name / Cell

Left BLANK, always — no data source, or a same-day production call Brian
makes by hand (2026-09-03: "the granular items we will add by hand"):
  - the whole SCHEDULE table (Crew Call through Load Out/Curfew — the
    minute-by-minute choreography is a day-of call, not form data)
  - Engineer (FOH/Mon names), Consoles, PA, Subs, LIGHTING (Pre-Scheduled/
    Live), VIDEO, Buyout

  python3 daysheet.py --event 1
"""
import argparse
import re
import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent
for _cand in (HERE.parent, HERE.parent / "app"):
    if (_cand / "advance_db.py").exists():
        sys.path.insert(0, str(_cand))
        break
sys.path.insert(0, str(HERE))
import advance_db as db

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from urllib.parse import quote

TEMPLATES = HERE / "doc_templates"
FILLED = HERE / "filled"
FILLED.mkdir(exist_ok=True)
TEMPLATE_BY_ACTS = {
    1: TEMPLATES / "FSQ Single Band Advance.docx",
    2: TEMPLATES / "FSQ 2 Band Advance.docx",
    3: TEMPLATES / "FSQ 3 Band Advance.docx",
}
DEFAULT_TEMPLATE = TEMPLATE_BY_ACTS[1]  # kept for callers that don't pass one


def norm(s):
    return re.sub(r"\s+", " ", (s or "").strip()).lower().rstrip("?")


def _yn(b):
    return "Yes" if b is True else ("No" if b is False else None)


def form_fields(sub):
    """Flatten a form submission to sheet-key -> string, for merging."""
    if not sub:
        return {}
    d = sub.get("data") or {}
    out = {
        "contact_name": sub.get("contact_name") or d.get("contact_name"),
        "contact_phone": sub.get("contact_phone") or d.get("contact_phone"),
        "stage_type": sub.get("stage_type") or d.get("stage_type"),
        "monitors": (str(sub["monitors"]) if sub.get("monitors") is not None
                     else d.get("monitors")),
        "own_iems": _yn(sub.get("own_iems")) or d.get("own_iems"),
        "split_snake": sub.get("split_snake") or d.get("split_snake"),
        "stage_plot_desc": d.get("stage_plot_desc"),
        "input_notes": d.get("input_notes"),
        "backline": d.get("backline"),
        "own_engineer": sub.get("own_engineer") or d.get("own_engineer"),
        "scenic": d.get("scenic"),
        "lighting": d.get("lighting"),
        "merch": _yn(sub.get("merch")) or d.get("merch"),
        "band_tent": sub.get("band_tent") or d.get("band_tent"),
        "performers": (str(sub["performers"]) if sub.get("performers") is not None
                       else d.get("performers")),
        "large_vehicle": _yn(sub.get("large_vehicle")) or d.get("large_vehicle"),
        "stage_plot_file": d.get("stage_plot_file"),
    }
    return {k: v for k, v in out.items() if v not in (None, "")}


def merged_fields(act):
    """Spreadsheet overrides win; the form fills the blanks."""
    f = form_fields(act.get("submission"))
    for k, v in (act.get("sheet_fields") or {}).items():
        if v not in (None, ""):
            f[k] = v
    return f


def act_row_values(f):
    """Merged fields -> {normalized day-sheet row label: cell text}, matching
    the current templates' row set (Monitors/IEMs and Stage Type/Scenic
    Notes are separate rows now, not blended)."""
    out = {}
    if not f:
        return out

    saved = f.get("_stageplot_saved")
    if saved:
        out["stage plot"] = f"See DB — {saved}"
    else:
        sp = f.get("stage_plot_desc", "")
        if f.get("stage_plot_file"):
            sp = (sp + " (file on record)").strip()
        if sp:
            out["stage plot"] = sp

    if f.get("monitors") not in (None, ""):
        out["monitors"] = f"{f['monitors']} wedges"
    if f.get("own_iems"):
        out["iems"] = checkbox_pair("Yes", "No", f["own_iems"])

    notes = []
    if f.get("input_notes"):
        notes.append(f["input_notes"])
    if str(f.get("own_iems", "")).lower() == "yes" and f.get("split_snake"):
        notes.append(f"Split snake: {f['split_snake']}")
    if f.get("lighting"):
        notes.append(f"Lighting request: {f['lighting']}")
    if notes:
        out["input notes"] = " · ".join(notes)

    if f.get("stage_type"):
        out["stage type"] = checkbox_pair(
            "Flat", "Riser", "Riser" if "riser" in str(f["stage_type"]).lower() else "Flat")
    if f.get("scenic"):
        out["scenic notes"] = f["scenic"]

    if f.get("merch"):
        out["merch"] = f["merch"]
    if f.get("large_vehicle"):
        out["parking"] = "Large vehicle" if str(f["large_vehicle"]).lower() == "yes" else "Standard"
    if f.get("performers") not in (None, ""):
        out["drink tix"] = str(f["performers"])
    if f.get("band_tent"):
        out["dressing room tent"] = "Yes" if str(f["band_tent"]).lower().startswith("yes") else "No"
    if f.get("backline"):
        out["backline"] = f["backline"]
    if f.get("contact_name"):
        out["band contact — name"] = f["contact_name"]
    if f.get("contact_phone"):
        out["band contact — cell"] = f["contact_phone"]
    return out


def checkbox_pair(label_a, label_b, chosen):
    """'☐ A     ☐ B' with whichever of A/B was chosen swapped to ☒."""
    a = "☒" if norm(chosen) == norm(label_a) else "☐"
    b = "☒" if norm(chosen) == norm(label_b) else "☐"
    return f"{a} {label_a}     {b} {label_b}"


def set_cell(cell, text):
    text = "" if text is None else str(text)
    p = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    for r in list(p.runs):
        r.text = ""
    (p.runs[0] if p.runs else p.add_run("")).text = text


def set_cell_link(cell, text, target):
    """Replace a cell's content with a single clickable hyperlink (blue, underlined).
    `target` is a relative path — the stage plot sits in the same folder as the doc,
    so Word resolves it wherever the folder lives."""
    p = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    r_id = cell.part.relate_to(target, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run)
    p._p.append(link)


def set_para_text(paragraph, text):
    for extra in list(paragraph.runs[1:]):
        extra._element.getparent().remove(extra._element)
    (paragraph.runs[0] if paragraph.runs else paragraph.add_run("")).text = text


def append_to_run(paragraph, prefix, value):
    """Find the run starting with `prefix` and set it to prefix+value — used
    for label+blank lines like 'TONIGHT — MC: ' that have no separate blank
    run of their own to target."""
    for r in paragraph.runs:
        if r.text.strip().startswith(prefix.strip()):
            r.text = f"{prefix}{value}"
            return True
    return False


def all_tables(doc):
    out = []
    def rec(tbls):
        for t in tbls:
            out.append(t)
            for row in t.rows:
                for cell in row.cells:
                    rec(cell.tables)
    rec(doc.tables)
    return out


def find_grid(doc):
    """The one big EVENT INFORMATION -> MISC table (right-hand side)."""
    for t in all_tables(doc):
        for r in t.rows:
            if r.cells and norm(r.cells[0].text) == "event information":
                return t
    return None


def find_lead_table(doc):
    """The small Lead/Cell table (left-hand side)."""
    for t in all_tables(doc):
        if len(t.rows) == 2 and norm(t.rows[0].cells[0].text) == "lead":
            return t
    return None


def fill_header(grid, event, acts, single):
    """EVENT INFORMATION value cell: Date / Event / Band (single-band only) /
    TONIGHT — MC/DJ (multi-band only) — each its own paragraph."""
    for r in grid.rows:
        if r.cells and norm(r.cells[0].text) == "event information":
            value_cell = r.cells[1]
            date = event.get("event_date").isoformat() if event.get("event_date") else ""
            for p in value_cell.paragraphs:
                t = p.text.strip()
                if t.startswith("Date:"):
                    set_para_text(p, f"Date: {date}" if date else "Date:")
                elif t.startswith("Event:"):
                    set_para_text(p, f"Event: {event.get('name') or ''}")
                elif t.startswith("Band:") and single:
                    bands = ", ".join(a["artist"]["name"] for a in acts if a.get("artist"))
                    set_para_text(p, f"Band: {bands}" if bands else "Band:")
                elif t.startswith("TONIGHT"):
                    det = event.get("details") or {}
                    if det.get("mc"):
                        append_to_run(p, "TONIGHT — MC: ", det["mc"])
                    if det.get("dj"):
                        append_to_run(p, "     DJ: ", det["dj"])
            return


def fill_event_type(grid, event):
    det = event.get("details") or {}
    for r in grid.rows:
        if r.cells and norm(r.cells[0].text) == "event type":
            value_cell = r.cells[1]
            paras = value_cell.paragraphs
            if len(paras) >= 1 and det.get("event_type"):
                set_para_text(paras[0], checkbox_pair(
                    "Internal Event:", "Third Party Event:",
                    "Internal Event:" if norm(det["event_type"]) == "internal" else "Third Party Event:",
                ))
            if len(paras) >= 2 and det.get("paying_band"):
                yn = "Yes" if norm(det["paying_band"]) == "yes" else "No"
                set_para_text(paras[1], f"Are we paying the band?   {checkbox_pair('Yes', 'No', yn)}")
            return


def fill_lead(doc, event):
    det = event.get("details") or {}
    t = find_lead_table(doc)
    if not t:
        return
    if det.get("lead_name"):
        set_cell(t.rows[0].cells[1], det["lead_name"])
    if det.get("lead_phone"):
        set_cell(t.rows[1].cells[1], det["lead_phone"])


def act_columns(grid, n_acts):
    """Row -> its N value cells, keyed by normalized label. Cells inside a
    colSpan report once per spanned grid column in python-docx, so a
    single-value row (Event Type, MISC header, …) still resolves fine —
    callers that expect N distinct act cells use cells[1:1+n_acts]."""
    rows = {}
    for r in grid.rows:
        label = norm(r.cells[0].text)
        if label:
            rows.setdefault(label, r)
    return rows


def fill(event_id, template=None, out_path=None, stageplot_names=None):
    """stageplot_names: {artist name -> filed stage-plot filename}. When an act's
    plot was downloaded and filed next to this doc, its Stage Plot cell reads
    'See DB — <filename>' instead of the inline description."""
    stageplot_names = stageplot_names or {}
    with db.get_conn() as conn, conn.cursor() as cur:
        event = db.get_event(cur, event_id)
        if not event:
            print(f"No event {event_id}", file=sys.stderr); sys.exit(1)
        acts = db.event_acts(cur, event_id)

    n = len(acts)
    if template is None:
        template = TEMPLATE_BY_ACTS.get(n)
        if template is None:
            print(f"Event {event_id} has {n} act(s) — no FSQ template for that "
                  f"count (1/2/3 only).", file=sys.stderr)
            sys.exit(1)
    if not template.exists():
        print(f"Template not found: {template}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(template))
    grid = find_grid(doc)
    if grid is None:
        print("Couldn't find the EVENT INFORMATION table.", file=sys.stderr); sys.exit(1)

    single = n == 1
    fill_header(grid, event, acts, single)
    fill_event_type(grid, event)
    fill_lead(doc, event)

    # multi-band act-name header row: bold slot label already printed by the
    # template, second paragraph is the blank line for the actual band name
    if not single:
        for r in grid.rows:
            if r.cells and r.cells[0].text.strip() == "" and any(
                    c.paragraphs and c.paragraphs[0].runs and c.paragraphs[0].runs[0].bold
                    for c in r.cells[1:1 + n]):
                for i, a in enumerate(acts):
                    if not a.get("artist"):
                        continue
                    ci = 1 + i
                    if ci < len(r.cells) and len(r.cells[ci].paragraphs) >= 2:
                        set_para_text(r.cells[ci].paragraphs[1], a["artist"]["name"])
                break

    rows_by_label = act_columns(grid, n)
    filled_acts = 0
    for i, a in enumerate(acts):
        mf = merged_fields(a)
        if a.get("set_time"):
            mf_set_length = a["set_time"]
        else:
            mf_set_length = None
        who = a["artist"]["name"] if a.get("artist") else None
        saved_plot = stageplot_names.get(who) if who else None
        if saved_plot:
            mf["_stageplot_saved"] = saved_plot
        values = act_row_values(mf)
        if mf_set_length:
            values["set length"] = mf_set_length
        if values:
            filled_acts += 1
        for label, text in values.items():
            row = rows_by_label.get(label)
            if not row:
                continue
            ci = 1 + i
            if ci >= len(row.cells):
                continue
            if label == "stage plot" and saved_plot:
                set_cell_link(row.cells[ci], text, quote(saved_plot))
            else:
                set_cell(row.cells[ci], text)

    if out_path is None:
        tag = re.sub(r"[^A-Za-z0-9]+", "_", event.get("name") or f"event{event_id}").strip("_")
        date = event.get("event_date")
        out_path = FILLED / f"{tag}__{date or 'nodate'}__daysheet.docx"
    doc.save(out_path)
    print(f"Filled day-sheet ({n}-band template): {out_path}")
    print(f"  {event.get('name')} @ {event.get('venue')} {event.get('event_date')} "
          f"— {filled_acts}/{len(acts)} act(s) filled")
    for a in acts:
        who = a["artist"]["name"] if a.get("artist") else "(none)"
        src = []
        if a.get("submission"):
            src.append("form")
        if a.get("sheet_fields"):
            src.append("sheet")
        tag = ("+".join(src)) if src else "no data"
        print(f"    {a['slot']:16} {who}  [{tag}]  {a.get('set_time') or ''}")
    return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--event", type=int, required=True)
    ap.add_argument("--template", type=Path, default=None,
                     help="override the auto-picked (by act count) template")
    ap.add_argument("--out", type=Path, default=None)
    args = ap.parse_args()
    fill(args.event, args.template, args.out)


if __name__ == "__main__":
    main()
